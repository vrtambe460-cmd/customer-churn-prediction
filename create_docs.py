"""Create properly formatted Word documents"""

import subprocess
import sys

subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color):
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)

def add_formatted_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, '4472C4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'D9E2F3')

    doc.add_paragraph()

def create_synopsis():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_heading('CUSTOMER CHURN PREDICTION SYSTEM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Using Multi-Agent Machine Learning Architecture')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    doc.add_paragraph()

    doc.add_heading('PROJECT SYNOPSIS', level=1)

    doc.add_heading('1. Project Title', level=2)
    doc.add_paragraph('Customer Churn Prediction System Using Multi-Agent Machine Learning Architecture')

    doc.add_heading('2. Objective', level=2)
    doc.add_paragraph('To build an intelligent system that predicts which telecom customers are likely to leave (churn) and provides actionable retention recommendations using a multi-agent machine learning architecture.')

    doc.add_heading('3. Problem Statement', level=2)
    doc.add_paragraph('Telecom companies lose significant revenue when customers leave. Identifying at-risk customers early allows companies to take proactive retention measures. This project automates the prediction process using AI agents.')

    doc.add_heading('4. Solution Overview', level=2)
    doc.add_paragraph('The project implements a 3-agent architecture:')
    doc.add_paragraph('DataPrepAgent - Handles data cleaning and preprocessing', style='List Bullet')
    doc.add_paragraph('PredictionAgent - Trains ML models and makes predictions', style='List Bullet')
    doc.add_paragraph('RecommendationAgent - Generates retention strategies', style='List Bullet')
    doc.add_paragraph('A Streamlit web dashboard provides an interactive interface for users to input customer data, view predictions, analyze trends, and upload bulk CSV files.')

    doc.add_heading('5. Technologies Used', level=2)
    add_formatted_table(doc,
        ['Category', 'Technology'],
        [
            ['Language', 'Python 3.8+'],
            ['Machine Learning', 'Scikit-learn, XGBoost'],
            ['Data Processing', 'Pandas, NumPy'],
            ['Web Framework', 'Streamlit'],
            ['Visualization', 'Plotly'],
            ['Model Persistence', 'Joblib'],
        ]
    )

    doc.add_heading('6. ML Models Implemented', level=2)
    add_formatted_table(doc,
        ['Model', 'Type', 'Accuracy'],
        [
            ['Logistic Regression', 'Linear', '79.48%'],
            ['Random Forest', 'Ensemble', '78.69%'],
            ['XGBoost', 'Gradient Boosting', '78.20%'],
        ]
    )

    doc.add_heading('7. Dataset', level=2)
    doc.add_paragraph('Source: Telco Customer Churn Dataset')
    doc.add_paragraph('Records: 5,043 customers')
    doc.add_paragraph('Features: 20 columns (demographics, services, billing, churn)')
    doc.add_paragraph('Target: Churn (Yes/No)')

    doc.add_heading('8. Features', level=2)
    doc.add_paragraph('Single customer prediction with risk gauge', style='List Bullet')
    doc.add_paragraph('Interactive data analytics charts', style='List Bullet')
    doc.add_paragraph('Model performance comparison', style='List Bullet')
    doc.add_paragraph('Bulk CSV upload and analysis', style='List Bullet')
    doc.add_paragraph('Downloadable results', style='List Bullet')

    doc.add_heading('9. Applications', level=2)
    doc.add_paragraph('Telecom customer retention', style='List Bullet')
    doc.add_paragraph('Subscription service management', style='List Bullet')
    doc.add_paragraph('SaaS churn reduction', style='List Bullet')
    doc.add_paragraph('Banking customer loyalty programs', style='List Bullet')

    doc.add_heading('10. Future Scope', level=2)
    doc.add_paragraph('Real-time API integration', style='List Bullet')
    doc.add_paragraph('Deep learning models (LSTM, Neural Networks)', style='List Bullet')
    doc.add_paragraph('Customer segmentation clustering', style='List Bullet')
    doc.add_paragraph('Automated email/SMS alerts', style='List Bullet')
    doc.add_paragraph('A/B testing for retention strategies', style='List Bullet')

    doc.add_heading('11. Conclusion', level=2)
    doc.add_paragraph('The system successfully predicts customer churn with ~79% accuracy using a modular multi-agent architecture, enabling data-driven retention decisions.')

    doc.add_heading('12. References', level=2)
    doc.add_paragraph('IBM Telco Customer Churn Dataset', style='List Bullet')
    doc.add_paragraph('Scikit-learn Documentation', style='List Bullet')
    doc.add_paragraph('Streamlit Documentation', style='List Bullet')
    doc.add_paragraph('XGBoost Documentation', style='List Bullet')

    doc.save(r'C:\Users\Lenovo\Desktop\Yuktamedia\ML project\1st\SYNOPSIS.docx')
    print('Created: SYNOPSIS.docx')

def create_documentation():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_heading('CUSTOMER CHURN PREDICTION SYSTEM', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Technical Documentation')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(68, 114, 196)

    doc.add_paragraph()

    doc.add_heading('TABLE OF CONTENTS', level=1)
    toc_items = [
        '1. Introduction',
        '2. System Architecture',
        '3. Installation Guide',
        '4. Project Structure',
        '5. Agent Documentation',
        '6. Usage Guide',
        '7. Dataset Description',
        '8. Model Details',
        '9. API Reference',
        '10. Troubleshooting',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    doc.add_heading('1. INTRODUCTION', level=1)

    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph('This system predicts customer churn in telecom companies using machine learning. It employs a multi-agent architecture where specialized agents handle different tasks.')

    doc.add_heading('1.2 Scope', level=2)
    doc.add_paragraph('Predicts if a customer will leave (churn)', style='List Bullet')
    doc.add_paragraph('Provides risk level assessment (Low/Medium/High)', style='List Bullet')
    doc.add_paragraph('Generates retention recommendations', style='List Bullet')
    doc.add_paragraph('Supports bulk analysis via CSV upload', style='List Bullet')

    doc.add_heading('1.3 Target Users', level=2)
    doc.add_paragraph('Telecom companies', style='List Bullet')
    doc.add_paragraph('Customer retention teams', style='List Bullet')
    doc.add_paragraph('Data analysts', style='List Bullet')
    doc.add_paragraph('Business managers', style='List Bullet')

    doc.add_page_break()

    doc.add_heading('2. SYSTEM ARCHITECTURE', level=1)

    doc.add_heading('2.1 High-Level Architecture', level=2)
    doc.add_paragraph('The system consists of three main components:')
    doc.add_paragraph('Streamlit Frontend - Web dashboard for user interaction', style='List Bullet')
    doc.add_paragraph('Main Orchestrator - Coordinates all agents', style='List Bullet')
    doc.add_paragraph('Three AI Agents - Handle specific tasks', style='List Bullet')

    doc.add_heading('2.2 Agent Architecture', level=2)
    add_formatted_table(doc,
        ['Agent', 'Responsibility', 'Input', 'Output'],
        [
            ['DataPrepAgent', 'Data cleaning', 'Raw CSV', 'Clean DataFrame'],
            ['PredictionAgent', 'ML training', 'Clean Data', 'Predictions'],
            ['RecommendationAgent', 'Advice generation', 'Predictions', 'Recommendations'],
        ]
    )

    doc.add_heading('2.3 Data Flow', level=2)
    doc.add_paragraph('User Input -> DataPrepAgent -> PredictionAgent -> RecommendationAgent -> Dashboard Output')

    doc.add_page_break()

    doc.add_heading('3. INSTALLATION GUIDE', level=1)

    doc.add_heading('3.1 Prerequisites', level=2)
    doc.add_paragraph('Python 3.8 or higher', style='List Bullet')
    doc.add_paragraph('pip package manager', style='List Bullet')
    doc.add_paragraph('Git (optional)', style='List Bullet')

    doc.add_heading('3.2 Install Steps', level=2)
    doc.add_paragraph('Step 1: Navigate to project directory')
    doc.add_paragraph('cd "C:\\Users\\Lenovo\\Desktop\\Yuktamedia\\ML project\\1st"')
    doc.add_paragraph('Step 2: Install dependencies')
    doc.add_paragraph('pip install -r requirements.txt')

    doc.add_heading('3.3 Required Packages', level=2)
    add_formatted_table(doc,
        ['Package', 'Version', 'Purpose'],
        [
            ['pandas', '>=1.5.0', 'Data manipulation'],
            ['numpy', '>=1.22.0', 'Numerical operations'],
            ['scikit-learn', '>=1.1.0', 'Machine learning'],
            ['xgboost', '>=1.6.0', 'Gradient boosting'],
            ['streamlit', '>=1.28.0', 'Web dashboard'],
            ['plotly', '>=5.15.0', 'Interactive charts'],
            ['joblib', '>=1.1.0', 'Model persistence'],
        ]
    )

    doc.add_page_break()

    doc.add_heading('4. PROJECT STRUCTURE', level=1)
    add_formatted_table(doc,
        ['File/Folder', 'Description'],
        [
            ['agents/', 'AI Agents package'],
            ['agents/base_agent.py', 'Abstract base class'],
            ['agents/data_prep_agent.py', 'Data preparation agent'],
            ['agents/prediction_agent.py', 'ML prediction agent'],
            ['agents/recommendation_agent.py', 'Recommendation agent'],
            ['models/', 'ML Models package'],
            ['data/', 'Dataset directory'],
            ['data/telco_churn.csv', 'Main dataset (5,043 records)'],
            ['utils/', 'Utilities package'],
            ['app.py', 'Streamlit web application'],
            ['main.py', 'Backend runner (CLI)'],
            ['requirements.txt', 'Python dependencies'],
        ]
    )

    doc.add_page_break()

    doc.add_heading('5. AGENT DOCUMENTATION', level=1)

    doc.add_heading('5.1 DataPrepAgent', level=2)
    doc.add_paragraph('Purpose: Handles all data preprocessing tasks.')
    doc.add_paragraph('Methods:')
    add_formatted_table(doc,
        ['Method', 'Description'],
        [
            ['initialize()', 'Setup agent resources'],
            ['execute(data)', 'Run full preprocessing pipeline'],
            ['_clean_data(df)', 'Remove duplicates, handle missing values'],
            ['_encode_features(df)', 'Convert categorical to numerical'],
            ['_scale_features(df)', 'Normalize numerical features'],
            ['prepare_single_customer(data)', 'Prepare single customer for prediction'],
        ]
    )

    doc.add_heading('5.2 PredictionAgent', level=2)
    doc.add_paragraph('Purpose: Trains ML models and makes predictions.')
    doc.add_paragraph('Methods:')
    add_formatted_table(doc,
        ['Method', 'Description'],
        [
            ['initialize()', 'Setup 3 ML models'],
            ['execute(data)', 'Train and evaluate models'],
            ['predict(X)', 'Make predictions on new data'],
            ['save_model(filepath)', 'Save best model to disk'],
            ['load_model(filepath)', 'Load model from disk'],
        ]
    )

    doc.add_heading('5.3 RecommendationAgent', level=2)
    doc.add_paragraph('Purpose: Generates customer retention recommendations.')
    doc.add_paragraph('Risk Levels:')
    add_formatted_table(doc,
        ['Probability', 'Level', 'Action'],
        [
            ['>= 0.7', 'HIGH', 'Urgent intervention'],
            ['>= 0.4', 'MEDIUM', 'Offer promotions'],
            ['< 0.4', 'LOW', 'Maintain service'],
        ]
    )

    doc.add_page_break()

    doc.add_heading('6. USAGE GUIDE', level=1)

    doc.add_heading('6.1 Run Web Dashboard', level=2)
    doc.add_paragraph('Command: streamlit run app.py')
    doc.add_paragraph('Open browser: http://localhost:8501')

    doc.add_heading('6.2 Dashboard Tabs', level=2)
    add_formatted_table(doc,
        ['Tab', 'Function'],
        [
            ['Predict', 'Single customer prediction'],
            ['Analytics', 'Data visualization'],
            ['Models', 'Model comparison'],
            ['Bulk Analysis', 'CSV upload & analysis'],
        ]
    )

    doc.add_heading('6.3 Upload CSV Format', level=2)
    doc.add_paragraph('Required columns:')
    doc.add_paragraph('customerID, gender, SeniorCitizen, Partner, Dependents, tenure, PhoneService, MultipleLines, InternetService, OnlineSecurity, OnlineBackup, DeviceProtection, TechSupport, StreamingTV, StreamingMovies, Contract, PaperlessBilling, PaymentMethod, MonthlyCharges, TotalCharges, Churn')

    doc.add_page_break()

    doc.add_heading('7. DATASET DESCRIPTION', level=1)

    doc.add_heading('7.1 Source', level=2)
    doc.add_paragraph('Telco Customer Churn Dataset (IBM Sample Data)')

    doc.add_heading('7.2 Size', level=2)
    doc.add_paragraph('Records: 5,043 customers')
    doc.add_paragraph('Features: 20 columns')
    doc.add_paragraph('Target: Churn (Yes/No)')

    doc.add_heading('7.3 Feature Descriptions', level=2)
    add_formatted_table(doc,
        ['Feature', 'Type', 'Description'],
        [
            ['gender', 'Categorical', 'Male/Female'],
            ['SeniorCitizen', 'Binary', '0/1'],
            ['Partner', 'Binary', 'Yes/No'],
            ['Dependents', 'Binary', 'Yes/No'],
            ['tenure', 'Numerical', 'Months as customer'],
            ['PhoneService', 'Binary', 'Yes/No'],
            ['InternetService', 'Categorical', 'DSL/Fiber optic/No'],
            ['Contract', 'Categorical', 'Month-to-month/One year/Two year'],
            ['MonthlyCharges', 'Numerical', 'Monthly bill amount'],
            ['TotalCharges', 'Numerical', 'Total amount paid'],
            ['Churn', 'Binary', 'Yes/No (target)'],
        ]
    )

    doc.add_page_break()

    doc.add_heading('8. MODEL DETAILS', level=1)

    doc.add_heading('8.1 Logistic Regression', level=2)
    add_formatted_table(doc,
        ['Metric', 'Value'],
        [
            ['Accuracy', '79.48%'],
            ['Precision', '62.40%'],
            ['Recall', '56.55%'],
            ['F1 Score', '59.33%'],
        ]
    )

    doc.add_heading('8.2 Random Forest', level=2)
    add_formatted_table(doc,
        ['Metric', 'Value'],
        [
            ['Accuracy', '78.69%'],
            ['Precision', '61.93%'],
            ['Recall', '50.56%'],
            ['F1 Score', '55.67%'],
        ]
    )

    doc.add_heading('8.3 XGBoost', level=2)
    add_formatted_table(doc,
        ['Metric', 'Value'],
        [
            ['Accuracy', '78.20%'],
            ['Precision', '60.09%'],
            ['Recall', '52.43%'],
            ['F1 Score', '56.00%'],
        ]
    )

    doc.add_page_break()

    doc.add_heading('9. API REFERENCE', level=1)

    doc.add_heading('9.1 DataPrepAgent', level=2)
    doc.add_paragraph('from agents import DataPrepAgent')
    doc.add_paragraph('agent = DataPrepAgent()')
    doc.add_paragraph('agent.initialize()')
    doc.add_paragraph('result = agent.execute({"dataframe": df, "fit_scaler": True})')

    doc.add_heading('9.2 PredictionAgent', level=2)
    doc.add_paragraph('from agents import PredictionAgent')
    doc.add_paragraph('agent = PredictionAgent()')
    doc.add_paragraph('agent.initialize()')
    doc.add_paragraph('result = agent.execute({"dataframe": clean_df})')
    doc.add_paragraph('predictions, probabilities = agent.predict(test_df)')

    doc.add_heading('9.3 RecommendationAgent', level=2)
    doc.add_paragraph('from agents import RecommendationAgent')
    doc.add_paragraph('agent = RecommendationAgent()')
    doc.add_paragraph('agent.initialize()')
    doc.add_paragraph('result = agent.execute({"predictions": pred, "probabilities": prob})')

    doc.add_page_break()

    doc.add_heading('10. TROUBLESHOOTING', level=1)

    add_formatted_table(doc,
        ['Error', 'Cause', 'Solution'],
        [
            ['ModuleNotFoundError', 'Missing package', 'pip install -r requirements.txt'],
            ['FileNotFoundError', 'Wrong path', 'Check data/telco_churn.csv exists'],
            ['ValueError: Feature mismatch', 'Churn column in upload', 'Remove Churn column from CSV'],
            ['ConvergenceWarning', 'Model didnt converge', 'Increase max_iter or scale data'],
        ]
    )

    doc.add_heading('Port Already in Use', level=2)
    doc.add_paragraph('Find process: netstat -ano | findstr :8501')
    doc.add_paragraph('Kill process: taskkill /PID <PID> /F')

    doc.save(r'C:\Users\Lenovo\Desktop\Yuktamedia\ML project\1st\DOCUMENTATION.docx')
    print('Created: DOCUMENTATION.docx')

if __name__ == "__main__":
    create_synopsis()
    create_documentation()
    print("\nDone! Both Word documents created.")
